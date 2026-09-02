from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from zipfile import BadZipFile, ZipFile

# Bump this whenever extraction logic changes for any format. The indexer
# uses it (alongside size/modified_at) to decide whether a previously
# indexed file needs to be re-extracted even though it hasn't changed on
# disk, so old content isn't silently kept forever after a parser fix.
CURRENT_EXTRACTOR_VERSION = 2

EXTRACTION_STATUSES = {"extracted", "unavailable", "skipped", "error", "timeout", "not_applicable"}


@dataclass(frozen=True, eq=False)
class ExtractionResult:
    text: str
    status: str
    detail: str = ""

    def __post_init__(self) -> None:
        if self.status not in EXTRACTION_STATUSES:
            raise ValueError(f"Unsupported extraction status: {self.status}")

    def __eq__(self, other: object) -> bool:
        if isinstance(other, ExtractionResult):
            return (self.text, self.status, self.detail) == (other.text, other.status, other.detail)
        if isinstance(other, str):
            return self.legacy_text == other
        return NotImplemented

    @property
    def legacy_text(self) -> str:
        if self.status == "extracted":
            return self.text
        prefix = {
            "unavailable": "extraction unavailable",
            "skipped": "extraction skipped",
            "error": "extraction error",
            "timeout": "extraction timed out",
        }.get(self.status, "extraction unavailable")
        return f"[{prefix}: {self.detail}]" if self.detail else f"[{prefix}]"

TEXT_EXTENSIONS = {
    ".txt", ".csv", ".tsv", ".md", ".log", ".json", ".xml", ".html", ".htm",
    ".yaml", ".yml", ".ini", ".cfg", ".conf", ".sql", ".py", ".js", ".ts",
    ".tsx", ".jsx", ".css", ".scss", ".java", ".c", ".h", ".cpp", ".hpp",
    ".cs", ".go", ".rs", ".php", ".sh", ".ps1", ".bat", ".properties",
}
OFFICE_XML_EXTENSIONS = {".docx", ".xlsx", ".pptx", ".odt", ".ods", ".odp"}
ARCHIVE_EXTENSIONS = {".zip"}
LEGACY_OFFICE_EXTENSIONS = {".doc", ".xls", ".ppt"}
OCR_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}


def extract_file(
    path: Path,
    max_chars: int,
    max_file_size_bytes: int = 512 * 1024 * 1024,
    enable_legacy_office: bool = False,
    libreoffice_command: str = "soffice",
    enable_ocr: bool = False,
    tesseract_command: str = "tesseract",
    ocrmypdf_command: str = "ocrmypdf",
    external_extraction_timeout_seconds: int = 120,
    external_extractors: dict[str, list[str]] | None = None,
) -> ExtractionResult:
    try:
        size = path.stat().st_size
        if size > max_file_size_bytes:
            return ExtractionResult("", "skipped", f"file exceeds {max_file_size_bytes} bytes")
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            result = _from_legacy(_extract_pdf(path, max_chars))
            if result.status != "unavailable" or not enable_ocr:
                return result
            if not _has_command(ocrmypdf_command):
                return ExtractionResult("", "unavailable", "OCRmyPDF is not installed")
            return ExtractionResult(
                _extract_with_ocrmypdf(path, max_chars, ocrmypdf_command, external_extraction_timeout_seconds),
                "extracted",
            )
        if suffix == ".docx":
            return _from_legacy(_extract_docx(path, max_chars))
        if suffix in OFFICE_XML_EXTENSIONS:
            return _from_legacy(_extract_zip_xml(path, max_chars))
        if suffix in ARCHIVE_EXTENSIONS:
            return _from_legacy(_extract_archive_manifest(path, max_chars))
        if suffix in LEGACY_OFFICE_EXTENSIONS:
            if not enable_legacy_office:
                return ExtractionResult("", "unavailable", "legacy Office extraction is disabled")
            if not _has_command(libreoffice_command):
                return ExtractionResult("", "unavailable", "LibreOffice is not installed")
            return ExtractionResult(
                _extract_with_libreoffice(path, max_chars, libreoffice_command, external_extraction_timeout_seconds),
                "extracted",
            )
        if suffix in OCR_IMAGE_EXTENSIONS:
            if not enable_ocr:
                return ExtractionResult("", "unavailable", "OCR is disabled")
            if not _has_command(tesseract_command):
                return ExtractionResult("", "unavailable", "Tesseract is not installed")
            return ExtractionResult(
                _extract_with_tesseract(path, max_chars, tesseract_command, external_extraction_timeout_seconds),
                "extracted",
            )
        if external_extractors and suffix in external_extractors:
            command = external_extractors[suffix]
            if not command or not _has_command(command[0]):
                return ExtractionResult("", "unavailable", f"parser is not installed for {suffix}")
            return ExtractionResult(
                _extract_with_plugin(path, max_chars, command, external_extraction_timeout_seconds),
                "extracted",
            )
        if suffix in TEXT_EXTENSIONS:
            return ExtractionResult(_extract_plain_text(path, max_chars), "extracted")
    except (OSError, BadZipFile, ValueError) as exc:
        return ExtractionResult("", "error", type(exc).__name__)
    except Exception as exc:
        return ExtractionResult("", "error", type(exc).__name__)
    return ExtractionResult("", "unavailable", "unsupported file type")


def extract_text(path: Path, max_chars: int, max_file_size_bytes: int = 512 * 1024 * 1024, **options: object) -> str:
    """Compatibility wrapper returning the historical marker format."""
    return extract_file(path, max_chars, max_file_size_bytes, **options).legacy_text


def _from_legacy(text: str) -> ExtractionResult:
    if text.startswith("[extraction unavailable:"):
        return ExtractionResult("", "unavailable", text[len("[extraction unavailable:") : -1].strip())
    return ExtractionResult(text, "extracted")


def _has_command(command: str) -> bool:
    return bool(shutil.which(command) or Path(command).is_file())


def _extract_with_plugin(path: Path, max_chars: int, command: list[str], timeout: int) -> str:
    completed = subprocess.run(
        [*command, str(path)],
        capture_output=True,
        text=True,
        timeout=timeout,
        check=False,
        shell=False,
    )
    if completed.returncode != 0:
        raise RuntimeError("configured parser returned an error")
    return completed.stdout[:max_chars]


def _extract_with_libreoffice(path: Path, max_chars: int, command: str, timeout: int) -> str:
    with tempfile.TemporaryDirectory(prefix="seamtech-office-") as output_dir:
        completed = subprocess.run(
            [command, "--headless", "--convert-to", "txt:Text", "--outdir", output_dir, str(path)],
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
        )
        output_path = Path(output_dir) / f"{path.stem}.txt"
        if completed.returncode != 0 or not output_path.exists():
            raise RuntimeError("LibreOffice could not convert the document")
        return output_path.read_text(encoding="utf-8", errors="ignore")[:max_chars]


def _extract_with_tesseract(path: Path, max_chars: int, command: str, timeout: int) -> str:
    completed = subprocess.run(
        [command, str(path), "stdout"],
        capture_output=True,
        text=True,
        timeout=timeout,
        check=False,
    )
    if completed.returncode != 0:
        raise RuntimeError("Tesseract could not read the image")
    return completed.stdout[:max_chars]


def _extract_with_ocrmypdf(path: Path, max_chars: int, command: str, timeout: int) -> str:
    with tempfile.TemporaryDirectory(prefix="seamtech-ocr-") as output_dir:
        output_path = Path(output_dir) / "extracted.txt"
        completed = subprocess.run(
            [command, "--sidecar", str(output_path), "--output-type", "pdf", str(path), str(Path(output_dir) / "output.pdf")],
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
        )
        if completed.returncode != 0 or not output_path.exists():
            raise RuntimeError("OCRmyPDF could not extract text from the PDF")
        return output_path.read_text(encoding="utf-8", errors="ignore")[:max_chars]


def _extract_pdf(path: Path, max_chars: int) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path), strict=False)
    chunks: list[str] = []
    total = 0
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            chunks.append(text)
            total += len(text)
        if total >= max_chars:
            break
    return "\n".join(chunks)[:max_chars] or "[extraction unavailable: PDF contains no embedded text]"


def _extract_docx(path: Path, max_chars: int) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        chunks.extend(cell.text for row in table.rows for cell in row.cells)
    text = "\n".join(chunks).strip()
    return text[:max_chars] or "[extraction unavailable: document contains no text]"


def _extract_zip_xml(path: Path, max_chars: int) -> str:
    chunks: list[str] = []
    with ZipFile(path) as archive:
        members = [info for info in archive.infolist() if not info.is_dir() and info.file_size <= 10 * 1024 * 1024]
        for info in members[:500]:
            if not info.filename.lower().endswith((".xml", ".rels")):
                continue
            raw = archive.read(info)
            text = re.sub(r"<[^>]+>", " ", raw.decode("utf-8", errors="ignore"))
            text = re.sub(r"\s+", " ", text).strip()
            if text:
                chunks.append(text)
            if sum(map(len, chunks)) >= max_chars:
                break
    return "\n".join(chunks)[:max_chars] or "[extraction unavailable: archive contains no readable text]"


def _extract_archive_manifest(path: Path, max_chars: int) -> str:
    with ZipFile(path) as archive:
        names = [info.filename for info in archive.infolist()[:10_000]]
    return "[archive members]\n" + "\n".join(names)[:max_chars]


def _extract_plain_text(path: Path, max_chars: int) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
